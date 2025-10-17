import markdown2
from pathlib import Path
from typing import List, Dict
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re


class ExportFormatter:
    """Handles exporting newsletters in multiple formats"""
    
    def __init__(self, output_dir: Path):
        self.output_dir = output_dir
    
    def create_html(self, markdown_content: str, output_filename: str = "newsletter.html") -> Path:
        """Convert markdown to HTML format for Notion/web browsers"""
        
        # Convert markdown to HTML with extras
        html_content = markdown2.markdown(
            markdown_content,
            extras=["fenced-code-blocks", "tables", "break-on-newline"]
        )
        
        # Create a styled HTML document
        html_template = f"""<!DOCTYPE html>
<html lang="sl">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Newsletter</title>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, 'Helvetica Neue', Arial, sans-serif;
            line-height: 1.8;
            max-width: 800px;
            margin: 0 auto;
            padding: 40px 20px;
            color: #333;
            background-color: #fff;
        }}
        h1 {{
            color: #1a1a1a;
            font-size: 2.5em;
            margin-bottom: 0.5em;
            line-height: 1.3;
            border-bottom: 3px solid #667eea;
            padding-bottom: 10px;
        }}
        h2 {{
            color: #2a2a2a;
            font-size: 2em;
            margin-top: 1.5em;
            margin-bottom: 0.5em;
        }}
        h3 {{
            color: #3a3a3a;
            font-size: 1.5em;
            margin-top: 1.2em;
            margin-bottom: 0.5em;
        }}
        p {{
            margin-bottom: 1.2em;
            font-size: 1.05em;
        }}
        img {{
            max-width: 100%;
            height: auto;
            border-radius: 8px;
            margin: 20px 0;
            box-shadow: 0 4px 12px rgba(0,0,0,0.1);
        }}
        em {{
            display: block;
            text-align: center;
            color: #666;
            font-size: 0.9em;
            margin-top: -10px;
            margin-bottom: 20px;
        }}
        strong {{
            color: #1a1a1a;
            font-weight: 600;
        }}
        ul, ol {{
            margin-left: 1.5em;
            margin-bottom: 1.2em;
        }}
        li {{
            margin-bottom: 0.5em;
        }}
        blockquote {{
            border-left: 4px solid #667eea;
            padding-left: 20px;
            margin: 20px 0;
            color: #555;
            font-style: italic;
        }}
        code {{
            background-color: #f5f5f5;
            padding: 2px 6px;
            border-radius: 3px;
            font-family: 'Courier New', monospace;
        }}
        pre {{
            background-color: #f5f5f5;
            padding: 15px;
            border-radius: 5px;
            overflow-x: auto;
        }}
    </style>
</head>
<body>
{html_content}
</body>
</html>"""
        
        # Write HTML file
        html_path = self.output_dir / output_filename
        html_path.write_text(html_template, encoding='utf-8')
        
        return html_path
    
    def create_docx(self, markdown_content: str, images_dir: Path, output_filename: str = "newsletter.docx") -> Path:
        """Convert markdown to DOCX format for Microsoft Word"""
        
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Parse markdown content
        lines = markdown_content.split('\n')
        
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                i += 1
                continue
            
            # Handle headings
            if line.startswith('# '):
                # H1 - Title
                heading = doc.add_heading(line[2:], level=1)
                heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in heading.runs:
                    run.font.color.rgb = RGBColor(26, 26, 26)
                    run.font.size = Pt(28)
            
            elif line.startswith('## '):
                # H2 - Section
                heading = doc.add_heading(line[3:], level=2)
                for run in heading.runs:
                    run.font.color.rgb = RGBColor(42, 42, 42)
                    run.font.size = Pt(22)
            
            elif line.startswith('### '):
                # H3 - Subsection
                heading = doc.add_heading(line[4:], level=3)
                for run in heading.runs:
                    run.font.color.rgb = RGBColor(58, 58, 58)
                    run.font.size = Pt(18)
            
            # Handle images
            elif line.startswith('!['):
                # Extract image path and alt text
                match = re.match(r'!\[(.*?)\]\((.*?)\)', line)
                if match:
                    alt_text = match.group(1)
                    image_path = match.group(2)
                    
                    # Construct full image path
                    full_image_path = self.output_dir / image_path
                    
                    if full_image_path.exists():
                        # Add image to document
                        doc.add_picture(str(full_image_path), width=Inches(6))
                        last_paragraph = doc.paragraphs[-1]
                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Add caption if next line is italic (caption)
                        if i + 1 < len(lines) and lines[i + 1].strip().startswith('*') and lines[i + 1].strip().endswith('*'):
                            caption_text = lines[i + 1].strip().strip('*')
                            caption = doc.add_paragraph(caption_text)
                            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in caption.runs:
                                run.font.italic = True
                                run.font.size = Pt(10)
                                run.font.color.rgb = RGBColor(102, 102, 102)
                            i += 1  # Skip caption line
            
            # Handle bold text
            elif '**' in line:
                paragraph = doc.add_paragraph()
                parts = re.split(r'(\*\*.*?\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = paragraph.add_run(part[2:-2])
                        run.bold = True
                    elif part:
                        paragraph.add_run(part)
            
            # Handle italic text
            elif line.startswith('*') and line.endswith('*') and not line.startswith('**'):
                paragraph = doc.add_paragraph(line.strip('*'))
                for run in paragraph.runs:
                    run.italic = True
                    run.font.color.rgb = RGBColor(102, 102, 102)
            
            # Regular paragraph
            else:
                # Remove markdown formatting for regular text
                clean_line = line.replace('**', '')
                if clean_line:
                    paragraph = doc.add_paragraph(clean_line)
                    paragraph.style = 'Normal'
                    for run in paragraph.runs:
                        run.font.size = Pt(11)
            
            i += 1
        
        # Save document
        docx_path = self.output_dir / output_filename
        doc.save(str(docx_path))
        
        return docx_path
    
    def create_all_formats(self, markdown_content: str, images_dir: Path) -> Dict[str, Path]:
        """Create all export formats"""
        
        formats = {}
        
        # Markdown (already exists)
        formats['markdown'] = self.output_dir / "newsletter.md"
        
        # HTML for Notion and web
        formats['html'] = self.create_html(markdown_content, "newsletter.html")
        
        # DOCX for Word
        formats['docx'] = self.create_docx(markdown_content, images_dir, "newsletter.docx")
        
        return formats

